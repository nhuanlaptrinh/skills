#!/usr/bin/env python3
"""Create a structured Vietnamese DOCX report from a small JSON document."""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell: Any, value: Any, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, table_data: dict[str, Any]) -> None:
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index in range(len(headers)):
            set_cell_text(cells[index], row[index] if index < len(row) else "")
    document.add_paragraph()


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(10.5)


def build_document(payload: dict[str, Any]) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(payload.get("title", "BÁO CÁO"))
    run.bold = True
    run.font.name = "Aptos Display"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = payload.get("subtitle")
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(11)

    for item in payload.get("metadata", []):
        add_paragraph(document, f"{item.get('label', '')}: {item.get('value', '')}")

    for section_data in payload.get("sections", []):
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(section_data.get("heading", ""))
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(31, 78, 121)
        for paragraph_text in section_data.get("paragraphs", []):
            add_paragraph(document, paragraph_text)
        for bullet in section_data.get("bullets", []):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(str(bullet))
            run.font.size = Pt(10.5)
        for table_data in section_data.get("tables", []):
            add_table(document, table_data)

    footer = document.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(payload.get("footer", "Báo cáo được lập từ dữ liệu đã cung cấp."))
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    return document


def validate_docx(path: Path) -> tuple[bool, str]:
    if not path.is_file() or path.stat().st_size < 2048:
        return False, "DOCX không tồn tại hoặc nhỏ bất thường"
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            required = {"[Content_Types].xml", "word/document.xml"}
            missing = required - names
            if missing:
                return False, f"Thiếu thành phần DOCX: {', '.join(sorted(missing))}"
            if archive.testzip() is not None:
                return False, "DOCX có mục ZIP hỏng"
    except (OSError, zipfile.BadZipFile) as exc:
        return False, f"Không đọc được DOCX: {exc}"
    return True, f"DOCX hợp lệ ({path.stat().st_size} bytes)"


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-json", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    try:
        payload = json.loads(args.input_json.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        print(f"input_error={exc}", file=sys.stderr)
        return 2
    if not isinstance(payload, dict):
        print("input_error=JSON root must be an object", file=sys.stderr)
        return 2
    if args.dry_run:
        print(json.dumps({"output": str(args.output), "sections": len(payload.get("sections", [])), "dry_run": True}, ensure_ascii=False))
        return 0
    args.output.parent.mkdir(parents=True, exist_ok=True)
    build_document(payload).save(args.output)
    valid, message = validate_docx(args.output)
    print(json.dumps({"output": str(args.output), "valid": valid, "message": message}, ensure_ascii=False))
    return 0 if valid else 1


if __name__ == "__main__":
    raise SystemExit(main())
